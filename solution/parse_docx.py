# parse_docx.py
"""
Parses the MDT proformas DOCX into a list of raw case dicts.
Each dict contains the plain-text content of the four logical sections
of one MDT proforma page, keyed for downstream extractors.

Output shape (one dict per case):
{
    "case_index":      int,          # 0-based
    "patient_details": str,          # raw text from Patient Details cell
    "staging":         str,          # raw text from Staging & Diagnosis cell
    "clinical":        str,          # raw text from Clinical Details cell
    "mdt_outcome":     str,          # raw text from MDT Outcome cell
    "mdt_date":        str | None,   # e.g. "07/03/2025" from the heading line
}
"""

import re
from pathlib import Path
from docx import Document
from docx.table import Table, _Cell

# ── helpers ───────────────────────────────────────────────────────────────────

def _cell_text(cell: _Cell) -> str:
    """Return all paragraph text in a cell, newline-separated, stripped."""
    return "\n".join(p.text for p in cell.paragraphs).strip()


def _table_section_text(table: Table, header_fragment: str) -> str | None:
    """
    Search every cell in `table` for one whose text starts with (or contains)
    `header_fragment`.  Return the *content* cells that follow in the same row,
    concatenated.  Returns None if not found.
    """
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if header_fragment.lower() in cell.text.lower():
                rest = [_cell_text(c) for c in row.cells[i + 1:] if _cell_text(c)]
                return "\n".join(rest) if rest else _cell_text(cell)
    return None


_DATE_RE = re.compile(
    r"Colorectal Multidisciplinary Meeting[\s\-\u2013\u2014]*(\d{1,2}/\d{1,2}/\d{4})",
    re.IGNORECASE,
)

def _extract_mdt_date(para_text: str) -> str | None:
    m = _DATE_RE.search(para_text)
    return m.group(1) if m else None


# ── section label fragments used to locate table rows ────────────────────────

_PATIENT_HEADER     = "Patient Details"
_STAGING_HEADER     = "Staging & Diagnosis"
_CLINICAL_HEADER    = "Clinical Details"
_MDT_OUTCOME_HEADER = "MDT Outcome"


def _extract_patient_block(table: Table) -> str:
    """
    Patient Details is the first data cell (row 0, col 0) of every proforma
    table.  It contains Hospital Number, NHS Number, Name, Gender, DOB.
    Row 0 is the header row; actual data lives in row 1, col 0.
    """
    try:
        return _cell_text(table.rows[1].cells[0])
    except IndexError:
        return ""


def _extract_section(table: Table, header: str) -> str:
    """
    Walks all rows; when a cell contains `header`, returns the text of
    the *next* row's first cell (the content row immediately below the header).
    Falls back to same-row remaining cells if no row follows.
    """
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            if header.lower() in cell.text.lower():
                if i + 1 < len(table.rows):
                    return _cell_text(table.rows[i + 1].cells[0])
                return _cell_text(cell)
    return ""


# ── main parser ───────────────────────────────────────────────────────────────

def _is_proforma_table(table: Table) -> bool:
    """
    Every MDT proforma page has a 3-column table whose first header cell
    contains the highlighted text 'Patient Details'.
    Skip tiny/malformed tables.
    """
    if not table.rows or len(table.rows) < 2:
        return False
    header_cell = _cell_text(table.rows[0].cells[0])
    return "Patient Details" in header_cell


def parse_cases(docx_path: str | Path) -> list[dict]:
    """
    Load `docx_path` and return a list of raw case dicts, one per proforma page.

    Each proforma occupies exactly one table in the document.  The MDT date is
    extracted from the paragraph that immediately precedes each table.
    """
    doc = Document(str(docx_path))

    # Build a flat list of document body elements preserving order.
    # python-docx exposes doc.paragraphs and doc.tables separately (not interleaved)
    # so we walk the underlying XML children directly to maintain para/table order.
    body = doc.element.body

    elements = []
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            from docx.text.paragraph import Paragraph
            elements.append(("para", Paragraph(child, doc)))
        elif tag == "tbl":
            from docx.table import Table as DocxTable
            elements.append(("table", DocxTable(child, doc)))

    cases: list[dict] = []
    last_para_text = ""

    for kind, elem in elements:
        if kind == "para":
            last_para_text = elem.text.strip()
        elif kind == "table":
            table: Table = elem
            if not _is_proforma_table(table):
                continue

            mdt_date = _extract_mdt_date(last_para_text)

            patient_details = _extract_patient_block(table)
            staging         = _extract_section(table, _STAGING_HEADER)
            clinical        = _extract_section(table, _CLINICAL_HEADER)
            mdt_outcome     = _extract_section(table, _MDT_OUTCOME_HEADER)

            cases.append({
                "case_index":      len(cases),
                "patient_details": patient_details,
                "staging":         staging,
                "clinical":        clinical,
                "mdt_outcome":     mdt_outcome,
                "mdt_date":        mdt_date,
            })

    return cases


# ── smoke test (run directly) ─────────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    path = sys.argv[1] if len(sys.argv) > 1 else "hackathon-mdt-outcome-proformas.docx"
    cases = parse_cases(path)
    print(f"Parsed {len(cases)} cases\n")
    for c in cases[:3]:
        print(f"=== Case {c['case_index']} | MDT date: {c['mdt_date']} ===")
        print("PATIENT:", c["patient_details"][:120])
        print("STAGING:", c["staging"][:120])
        print("CLINICAL:", c["clinical"][:120])
        print("OUTCOME:", c["mdt_outcome"][:120])
        print()
